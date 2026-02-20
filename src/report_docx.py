from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.analysis import compute_stats
from src.docx_chart import CHART_PLACEHOLDER_PREFIX, inject_editable_charts
from src.llm_client import generate_summary, infer_metric_unit

# 统计项显示：中文 (英文)
STAT_LABELS: Dict[str, str] = {
    "count": "计数 (count)",
    "mean": "平均值 (mean)",
    "min": "最小值 (min)",
    "max": "最大值 (max)",
    "start": "期初 (start)",
    "end": "期末 (end)",
    "abs_change": "绝对变化 (abs_change)",
    "pct_change": "变化率 (pct_change)",
}

# 报告字体：英文 Times New Roman，汉字 宋体
FONT_LATIN = "Times New Roman"
FONT_EAST_ASIA = "宋体"


def _set_run_fonts(run) -> None:
    """将 run 设为西文 Times New Roman、东亚 宋体。"""
    run.font.name = FONT_LATIN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def _apply_document_fonts(document: Document) -> None:
    """对文档中所有段落和表格单元格的 run 应用中英文字体。"""
    for para in document.paragraphs:
        for run in para.runs:
            _set_run_fonts(run)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_fonts(run)


def _format_number(value: Any) -> str:
    try:
        if value is None:
            return "-"
        if isinstance(value, float) and (value != value):
            return "-"
        return f"{value:.4f}"
    except Exception:
        return str(value)


def _sanitize_filename(text: str) -> str:
    cleaned = "".join(ch if ch not in "<>:\\/?*|\"" else "_" for ch in text)
    cleaned = cleaned.strip().replace("..", ".")
    return cleaned or "chart"


def _chart_categories_and_values(
    series_df: pd.DataFrame,
    date_col: str,
    value_col: str,
) -> Tuple[List[str], List[float]]:
    """根据时间跨度决定横轴按「月」或「日」：约一年按月，约一月按日。"""
    if series_df.empty:
        return [], []
    df = series_df[[date_col, value_col]].dropna()
    if df.empty:
        return [], []
    df = df.copy()
    df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
    df = df.dropna(subset=[date_col])
    if df.empty:
        return [], []
    span_days = (df[date_col].max() - df[date_col].min()).days
    if span_days > 60:
        df["_month"] = df[date_col].dt.to_period("M").astype(str)
        agg = df.groupby("_month", as_index=False).agg({value_col: "mean"})
        categories = agg["_month"].tolist()
        values = agg[value_col].tolist()
    else:
        df = df.sort_values(date_col)
        categories = df[date_col].dt.strftime("%m-%d").tolist()
        values = df[value_col].tolist()
    return categories, values


def _add_metrics_section(
    document: Document,
    date_range: str,
    metrics: List[str],
    df,
    date_col: str,
    config_path: str,
    units: Dict[str, str],
    chart_start_index: int,
) -> List[Tuple[List[str], List[float], str, Optional[str]]]:
    """向 document 追加「日期范围 + 各指标表格/占位符/结论」，返回本节的 chart_data。"""
    chart_data: List[Tuple[List[str], List[float], str, Optional[str]]] = []
    document.add_paragraph(f"日期范围: {date_range}")
    document.add_paragraph(f"指标数量: {len(metrics)}")

    for i, metric in enumerate(metrics):
        document.add_heading(metric, level=2)
        series_df = df[[date_col, metric]].dropna()
        stats = compute_stats(series_df[metric])

        table = document.add_table(rows=1, cols=2)
        table.style = "Light Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "统计项"
        hdr_cells[1].text = "值"
        for key, value in stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = STAT_LABELS.get(key, f"{key} ({key})")
            row_cells[1].text = _format_number(value)

        placeholder_idx = chart_start_index + i
        document.add_paragraph(f"{CHART_PLACEHOLDER_PREFIX}{placeholder_idx}")
        categories, vals = _chart_categories_and_values(series_df, date_col, metric)
        excel_unit = units.get(metric) or None
        try:
            unit = infer_metric_unit(config_path, str(metric), excel_unit=excel_unit)
        except Exception:
            unit = excel_unit or ""
        chart_data.append((categories, vals, str(metric), unit or None))

        try:
            summary = generate_summary(config_path, metric, stats, date_range)
            document.add_paragraph(summary)
        except Exception as exc:
            document.add_paragraph(f"结论生成失败: {exc}")

    return chart_data


def build_report(
    output_dir: str,
    title: str,
    date_range: str,
    metrics: List[str],
    df,
    date_col: str,
    config_path: str,
    units: Dict[str, str],
    output_path: Optional[str] = None,
) -> Tuple[str, List[Tuple[List[str], List[float], str, Optional[str]]]]:
    """生成新报告，返回 (docx 路径, chart_data)。若提供 output_path 则直接写入该路径（用于多轮共用同一文件）。"""
    os.makedirs(output_dir, exist_ok=True)
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{timestamp}.docx"
        output_path = os.path.join(output_dir, filename)

    document = Document()
    document.add_heading(title, level=1)
    chart_data = _add_metrics_section(
        document, date_range, metrics, df, date_col, config_path, units, chart_start_index=0
    )

    _apply_document_fonts(document)
    document.save(output_path)
    if chart_data:
        try:
            inject_editable_charts(output_path, chart_data)
        except Exception:
            pass
    return output_path, chart_data


def append_report_section(
    doc_path: str,
    section_title: str,
    date_range: str,
    metrics: List[str],
    df,
    date_col: str,
    config_path: str,
    units: Dict[str, str],
    chart_start_index: int,
) -> List[Tuple[List[str], List[float], str, Optional[str]]]:
    """向已有 docx 追加一节（多轮对话的一轮），占位符从 chart_start_index 起。返回本节 chart_data。"""
    document = Document(doc_path)
    document.add_heading(section_title, level=1)
    chart_data = _add_metrics_section(
        document, date_range, metrics, df, date_col, config_path, units, chart_start_index
    )
    _apply_document_fonts(document)
    document.save(doc_path)
    return chart_data


def append_summary_section(doc_path: str, title: str, content: str) -> None:
    """向已有 docx 末尾追加一节综合总结（仅标题与段落，无图表）。"""
    document = Document(doc_path)
    document.add_heading(title, level=1)
    for block in (content or "").strip().split("\n\n"):
        block = block.strip()
        if block:
            document.add_paragraph(block)
    _apply_document_fonts(document)
    document.save(doc_path)


def _iter_block_items(parent):
    """按文档顺序 yield 段落与表格（用于提取报告正文）。"""
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_report_text_summary(doc_path: str, stop_at_heading: Optional[str] = "综合总结") -> str:
    """从已生成的报告 docx 中提取正文（标题与表格统计），供综合总结 LLM 使用。
    读到标题 stop_at_heading 时停止，不包含该节及之后内容（避免把旧总结当数据）。"""
    if not os.path.isfile(doc_path):
        return ""
    try:
        document = Document(doc_path)
    except Exception:
        return ""
    lines: List[str] = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = (block.text or "").strip()
            if not text or text.startswith(CHART_PLACEHOLDER_PREFIX):
                continue
            if stop_at_heading and text.strip() == stop_at_heading:
                break
            lines.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                cells = [cell.text.strip() if cell.text else "" for cell in row.cells]
                if any(cells):
                    lines.append("\t".join(cells))
            lines.append("")
    return "\n".join(lines).strip()


def replace_summary_section(doc_path: str, title: str, content: str) -> None:
    """找到文档中已存在的「综合总结」节（按标题匹配），删除该节及其后所有内容，再追加新的标题与正文。"""
    document = Document(doc_path)
    body = document.element.body
    start_el = None
    for para in document.paragraphs:
        if (para.text or "").strip() == title.strip():
            start_el = para._element
            break
    if start_el is None:
        append_summary_section(doc_path, title, content)
        return
    children = list(body)
    try:
        start_idx = children.index(start_el)
    except ValueError:
        append_summary_section(doc_path, title, content)
        return
    # 保留末尾的 sectPr（节属性）
    end_idx = len(children)
    if children and "sectPr" in (children[-1].tag or ""):
        end_idx = len(children) - 1
    for i in range(start_idx, end_idx):
        body.remove(children[i])
    document.add_heading(title, level=1)
    for block in (content or "").strip().split("\n\n"):
        block = block.strip()
        if block:
            document.add_paragraph(block)
    _apply_document_fonts(document)
    document.save(doc_path)
